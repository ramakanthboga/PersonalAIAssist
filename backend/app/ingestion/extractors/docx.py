"""DOCX text extraction using python-docx."""

from __future__ import annotations

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.ingestion.extractors import ExtractedPage


def extract_docx(file_path: str) -> list[ExtractedPage]:
    """Extract text from a DOCX file. Treats entire document as one page."""
    try:
        doc = DocxDocument(file_path)
    except PackageNotFoundError as exc:
        raise ValueError(
            "Invalid or unsupported Word document. Only .docx (Office Open XML) is supported."
        ) from exc
    except Exception as exc:
        # Legacy binary .doc and corrupt ZIP packages surface as various exceptions.
        raise ValueError(f"Failed to read DOCX file: {exc}") from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)
    if not full_text.strip():
        return []

    return [ExtractedPage(
        text=full_text,
        page_number=1,
        metadata={"source_type": "docx", "paragraph_count": str(len(paragraphs))},
    )]
